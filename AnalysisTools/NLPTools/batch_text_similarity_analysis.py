# #!/usr/bin/env python
# # coding: utf-8
#
# # ## Word Embedding
#
#
# import os
# import numpy as np
# from docx import Document
# from openai import OpenAI
# from dotenv import load_dotenv
#
# load_dotenv()
# client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
#
#
# def get_embedding(text, model="text-embedding-3-small"):
#     text = text.replace("\n", " ")
#     return client.embeddings.create(input=[text], model=model).data[0].embedding
#
#
# def pad_embeddings(a, b):
#     target_length = max(len(a), len(b))
#     a = np.pad(a, (0, target_length - len(a)), 'constant')
#     b = np.pad(b, (0, target_length - len(b)), 'constant')
#     return a, b
#
#
# def cosine_similarity(a, b):
#     return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))
#
#
# baseline_cookie_text_en = """
# In the center of the scene, there's a woman standing at the sink doing the dishes. She has short hair and is wearing a sleeveless dress. Interestingly, water is overflowing from the sink and spilling onto the floor, but she seems quite calm and oblivious to the mess. Maybe she's lost in thought or just really focused on her task. To the left, there's a boy standing on a stool, reaching into a cupboard to grab some cookies from a jar. He seems quite determined and focused on getting those cookies. Below him, a girl is eagerly reaching up, hoping to get a cookie for herself. She's standing on the floor, looking excited and maybe a bit impatient. Through the window above the sink, there is a view of a peaceful landscape with trees and a fence. This adds a nice, tranquil contrast to the somewhat chaotic kitchen scene. The counter next to the sink has several plates and bowls. The overall atmosphere of the picture feels a bit chaotic but still homely. The woman at the sink seems to be in her own world, possibly relaxed or maybe daydreaming. The boy and girl, on the other hand, are full of energy and excitement, driven by their desire to get to the cookies. It seems like the kids are feeling mischievous and adventurous, while the woman is either very calm or completely unaware of the situation around her. This contrast creates a humorous and lively scene
# ."""
#
# baseline_cookie_text_es = """
# En el centro de la escena, hay una mujer parada en el fregadero lavando los platos. Tiene el cabello corto y lleva un vestido sin mangas. Curiosamente, el agua se está desbordando del fregadero y derramándose en el piso, pero ella parece bastante tranquila y ajena al desorden. Tal vez está perdida en sus pensamientos o simplemente muy concentrada en su tarea.
# A la izquierda, hay un niño parado en un taburete, alcanzando un frasco de galletas en el armario. Parece muy decidido y concentrado en conseguir esas galletas. Debajo de él, una niña está extendiendo la mano, esperando recibir una galleta para ella misma. Está parada en el suelo, luciendo emocionada y tal vez un poco impaciente.
# A través de la ventana encima del fregadero, hay una vista de un paisaje tranquilo con árboles y una cerca. Esto añade un contraste agradable y sereno a la escena algo caótica de la cocina.
# El mostrador al lado del fregadero tiene varios platos y tazones.
# La atmósfera general de la imagen se siente un poco caótica pero aún hogareña. La mujer en el fregadero parece estar en su propio mundo, posiblemente relajada o tal vez soñando despierta. Por otro lado, el niño y la niña están llenos de energía y emoción, impulsados por su deseo de alcanzar las galletas.
# Parece que los niños se sienten traviesos y aventureros, mientras que la mujer está muy calmada o completamente inconsciente de la situación a su alrededor. Este contraste crea una escena humorística y animada.
# """
#
# baseline_picnic_text_en = """
# First, right in the center, there’s a couple having a picnic. The man is sitting on the blanket, wearing glasses and sandals, and seems pretty absorbed in reading a book. The woman next to him is pouring something from a thermos into a cup, maybe coffee or tea, and she’s also sitting comfortably on the blanket. They seem relaxed and are enjoying the peaceful environment.
# Nearby, there’s a boy who looks very excited. He’s running towards the couple with a big smile on his face, and he’s holding a kite in one hand. The kite is flying high in the sky. There's also a dog running alongside the boy.
# To the right, closer to the water, there’s a young child playing on the beach. The child is building something in the sand, maybe a sandcastle, and they seem pretty focused on what they’re doing.
# In the background, on a dock by the water, there’s another person fishing. They’re standing at the edge of the dock, holding a fishing rod, and it looks like they might be waiting patiently for a catch. Further out on the water, you can see a sailboat with the number “470” on it, calmly sailing by.
# There are also a few other details in the picture, like a car parked in front of the house, a radio on the picnic blanket, and a flagpole near the water
# It’s like this calm, happy scene with everyone doing their own thing and just enjoying the day.
# ."""
#
# baseline_picnic_text_es = """
# En el centro, hay una pareja descansando sobre una manta. El hombre está súper concentrado leyendo un libro, como si no le importara nada más alrededor. Lleva gafas y está muy cómodo con sus sandalias. La mujer que está a su lado está sirviendo algo de un termo, tal vez café o té, y se ve que está relajada también.
# Luego, hay un niño que está corriendo hacia ellos, y se le ve muy contento. Tiene un paplote en la manoy hay un perro corriendo a su lado.
# Más a la derecha, cerca del agua, hay un niño pequeño jugando en la arena. Parece que está construyendo un castillo de arena o algo así, y está súper concentrado en lo que hace.
# En el fondo, se puede ver a alguien pescando desde un muelle. Está de pie, esperando a ver si pesca algo, imagino. Y más allá, en el agua, hay un velero con el número "470", que va navegando tranquilamente.
# También hay algunos otros detalles en la imagen, como un carro estacionado frente a la casa, una radio sobre la manta del picnic, y una bandera ondeando cerca del agua.
# En general, parece una escena calmada y feliz, con todos disfrutando del día a su manera.
# ."""
#
#
# def read_docx(file_path):
#     doc = Document(file_path)
#     full_text = []
#     for para in doc.paragraphs:
#         full_text.append(para.text)
#     return full_text
#
#
# def read_and_compute(file_path, baseline):
#     full_text_list = read_docx(file_path)
#
#     t1_start_index = 0
#     t1_end_index = 0
#     print(full_text_list)
#     for paragraph in full_text_list:
#         if ("storytelling" in paragraph.lower()) or ("mirar la imagen" in paragraph.lower()) or ("mira la imagen" in paragraph.lower()):
#             break
#         # This is for P16_S1, it doesn't have storytelling keyword.
#         if " You will look at the screen. Um you will see a picture" in paragraph:
#             break
#         t1_start_index += 1
#     for paragraph in full_text_list:
#         if "t2" in paragraph.lower():
#             break
#         t1_end_index += 1
#
#     t2_start_index = 0
#     t2_end_index = len(full_text_list) - 1
#     for paragraph in full_text_list:
#         if ("different game" in paragraph.lower()) or ("juego diferente" in paragraph.lower()) or ("animal" in paragraph.lower()) or ("frutas" in paragraph.lower()):
#             break
#         if "different challenge" in paragraph.lower():
#             break
#         t2_start_index += 1
#
#     t1 = (full_text_list[t1_start_index:t1_end_index])
#     t2 = (full_text_list[t2_start_index:t2_end_index])
#
#     t1_processed = list(
#         filter(lambda x: not ((x == "") or ("speaker 0" in x.lower()) or (x is None) or ("first prompt" in x.lower())),
#                t1))
#     t2_processed = list(
#         filter(lambda x: not ((x == "") or ("speaker 0" in x.lower()) or (x is None) or ("first prompt" in x.lower())),
#                t2))
#
#     for i in range(0, len(t1_processed)):
#         t1_processed[i] = t1_processed[i].replace("Speaker 1: ", "")
#
#     for i in range(0, len(t2_processed)):
#         t2_processed[i] = t2_processed[i].replace("Speaker 1: ", "")
#
#     test_text = "".join(t1_processed)
#
#     baseline_text_emb = get_embedding(baseline)
#     try:
#         test_text_emb = get_embedding(test_text)
#     except Exception as e:
#         print(f"Error: {e}")
#         print(f"Error in file: {file_path}")
#         print(f"Error in text: {test_text}")
#         print(f"Info: {t1_start_index}, {t1_end_index}, {t2_start_index}, {t2_end_index}")
#         return 0, "".join(t1_processed), "".join(t2_processed)
#     baseline_text_emb, test_text_emb = pad_embeddings(baseline_text_emb, test_text_emb)
#
#     return cosine_similarity(baseline_text_emb, test_text_emb), "".join(t1_processed), "".join(t2_processed)
#
#
# directory_path = "./data/ES/"
# output_path = "./data_processed/ES/"
#
# baseline_picnic_text = baseline_picnic_text_es
# baseline_cookie_text = baseline_cookie_text_es
#
# # Get all file names in the directory
# file_names = [f for f in os.listdir(directory_path) if os.path.isfile(os.path.join(directory_path, f))]
#
# with open(f"{output_path}similarity.txt", "w") as f:
#     f.write("Similarity Score: \n")
#
# # Print all file names
# for file_name in file_names:
#     print(f"Processing: {file_name}")
#     base_name = file_name.split(".")[0]
#     base_name.replace("_check", "")
#     if "S2" in base_name or "S4" in base_name:
#         baseline_text = baseline_picnic_text
#     else:
#         baseline_text = baseline_cookie_text
#     similarity, t1, t2 = read_and_compute(f"{directory_path}{file_name}", baseline_text)
#     print(f"{base_name} Similarity: {similarity:.2f}")
#     with open("similarity.txt", "a") as f:
#         f.write(f"{base_name}: {similarity:.2f}\n")
#
#     print(f"Writing processed T1 text to file: {output_path}{base_name}_t1.txt")
#     with open(f"{output_path}{base_name}_t1.txt", "w") as f:
#         f.write(t1)
#
#     print(f"Writing processed T2 text to file: {output_path}{base_name}_t2.txt")
#     with open(f"{output_path}{base_name}_t2.txt", "w") as f:
#         f.write(t2)

# !/usr/bin/env python
# coding: utf-8

# ## Word Embedding


import os
import re
import string
import numpy as np
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv

load_dotenv()
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))


def get_embedding(text, model="text-embedding-3-small"):
    text = text.replace("\n", " ")
    return client.embeddings.create(input=[text], model=model).data[0].embedding


def pad_embeddings(a, b):
    target_length = max(len(a), len(b))
    a = np.pad(a, (0, target_length - len(a)), 'constant')
    b = np.pad(b, (0, target_length - len(b)), 'constant')
    return a, b


def cosine_similarity(a, b):
    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))


# TODO: should I remove /n?

baseline_cookie_text_en = """
In the center of the scene, there's a woman standing at the sink doing the dishes. She has short hair and is wearing a sleeveless dress. Interestingly, water is overflowing from the sink and spilling onto the floor, but she seems quite calm and oblivious to the mess. Maybe she's lost in thought or just really focused on her task. To the left, there's a boy standing on a stool, reaching into a cupboard to grab some cookies from a jar. He seems quite determined and focused on getting those cookies. Below him, a girl is eagerly reaching up, hoping to get a cookie for herself. She's standing on the floor, looking excited and maybe a bit impatient. Through the window above the sink, there is a view of a peaceful landscape with trees and a fence. This adds a nice, tranquil contrast to the somewhat chaotic kitchen scene. The counter next to the sink has several plates and bowls. The overall atmosphere of the picture feels a bit chaotic but still homely. The woman at the sink seems to be in her own world, possibly relaxed or maybe daydreaming. The boy and girl, on the other hand, are full of energy and excitement, driven by their desire to get to the cookies. It seems like the kids are feeling mischievous and adventurous, while the woman is either very calm or completely unaware of the situation around her. This contrast creates a humorous and lively scene
."""

baseline_cookie_text_es = """
En el centro de la escena, hay una mujer parada en el fregadero lavando los platos. Tiene el cabello corto y lleva un vestido sin mangas. Curiosamente, el agua se está desbordando del fregadero y derramándose en el piso, pero ella parece bastante tranquila y ajena al desorden. Tal vez está perdida en sus pensamientos o simplemente muy concentrada en su tarea.
A la izquierda, hay un niño parado en un taburete, alcanzando un frasco de galletas en el armario. Parece muy decidido y concentrado en conseguir esas galletas. Debajo de él, una niña está extendiendo la mano, esperando recibir una galleta para ella misma. Está parada en el suelo, luciendo emocionada y tal vez un poco impaciente.
A través de la ventana encima del fregadero, hay una vista de un paisaje tranquilo con árboles y una cerca. Esto añade un contraste agradable y sereno a la escena algo caótica de la cocina.
El mostrador al lado del fregadero tiene varios platos y tazones.
La atmósfera general de la imagen se siente un poco caótica pero aún hogareña. La mujer en el fregadero parece estar en su propio mundo, posiblemente relajada o tal vez soñando despierta. Por otro lado, el niño y la niña están llenos de energía y emoción, impulsados por su deseo de alcanzar las galletas.
Parece que los niños se sienten traviesos y aventureros, mientras que la mujer está muy calmada o completamente inconsciente de la situación a su alrededor. Este contraste crea una escena humorística y animada.
"""

baseline_picnic_text_en = """
First, right in the center, there’s a couple having a picnic. The man is sitting on the blanket, wearing glasses and sandals, and seems pretty absorbed in reading a book. The woman next to him is pouring something from a thermos into a cup, maybe coffee or tea, and she’s also sitting comfortably on the blanket. They seem relaxed and are enjoying the peaceful environment. 
Nearby, there’s a boy who looks very excited. He’s running towards the couple with a big smile on his face, and he’s holding a kite in one hand. The kite is flying high in the sky. There's also a dog running alongside the boy.
To the right, closer to the water, there’s a young child playing on the beach. The child is building something in the sand, maybe a sandcastle, and they seem pretty focused on what they’re doing.
In the background, on a dock by the water, there’s another person fishing. They’re standing at the edge of the dock, holding a fishing rod, and it looks like they might be waiting patiently for a catch. Further out on the water, you can see a sailboat with the number “470” on it, calmly sailing by.
There are also a few other details in the picture, like a car parked in front of the house, a radio on the picnic blanket, and a flagpole near the water
It’s like this calm, happy scene with everyone doing their own thing and just enjoying the day. 
."""

baseline_picnic_text_es = """
En el centro, hay una pareja descansando sobre una manta. El hombre está súper concentrado leyendo un libro, como si no le importara nada más alrededor. Lleva gafas y está muy cómodo con sus sandalias. La mujer que está a su lado está sirviendo algo de un termo, tal vez café o té, y se ve que está relajada también.
Luego, hay un niño que está corriendo hacia ellos, y se le ve muy contento. Tiene un paplote en la manoy hay un perro corriendo a su lado.
Más a la derecha, cerca del agua, hay un niño pequeño jugando en la arena. Parece que está construyendo un castillo de arena o algo así, y está súper concentrado en lo que hace.
En el fondo, se puede ver a alguien pescando desde un muelle. Está de pie, esperando a ver si pesca algo, imagino. Y más allá, en el agua, hay un velero con el número "470", que va navegando tranquilamente.
También hay algunos otros detalles en la imagen, como un carro estacionado frente a la casa, una radio sobre la manta del picnic, y una bandera ondeando cerca del agua.
En general, parece una escena calmada y feliz, con todos disfrutando del día a su manera.
."""


def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return full_text


# def remove_punctuation(text):
#     translator = str.maketrans('', '', string.punctuation)
#     return text.translate(translator)

def remove_punctuation(text):
    text = re.sub(r'[^\w\s]', ' ', text)
    return re.sub(r'\s+', ' ', text)


def read_and_preprocess_text(file_path):
    full_text_list = read_docx(file_path)

    t1_start_index = 0
    t1_end_index = 0
    free_attempt_end_index = 0
    print(full_text_list)
    for paragraph in full_text_list:
        if ("storytelling" in paragraph.lower()) or ("mirar la imagen" in paragraph.lower()) or (
                "mira la imagen" in paragraph.lower()) or ("observa la imagen" in paragraph.lower()) or (
                "mires la imagen" in paragraph.lower()) or ("la imagen de esta pantalla" in paragraph.lower()) or (
                "una imagen en la pantalla" in paragraph.lower()) or ("Mirarás la imagen en " in paragraph):
            break
        # This is for P16_S1, it doesn't have storytelling keyword.
        if " You will look at the screen. Um you will see a picture" in paragraph:
            break
        t1_start_index += 1
    for paragraph in full_text_list:
        if "t2" in paragraph.lower():
            break
        t1_end_index += 1
    for paragraph in full_text_list:
        if "first prompt" in paragraph.lower():
            break
        free_attempt_end_index += 1

    t2_start_index = 0
    t2_end_index = len(full_text_list) - 1
    for paragraph in full_text_list:
        if ("different game" in paragraph.lower()) or ("juego diferente" in paragraph.lower()) or (
                "animal" in paragraph.lower()) or ("frutas" in paragraph.lower()):
            break
        if "different challenge" in paragraph.lower():
            break
        t2_start_index += 1

    # if t1_end_index > t2_start_index or t2_start_index > t2_end_index:
    #     print(f"Error in file: {file_path}")
    #     print(f"Info: {t1_start_index}, {t1_end_index}, {t2_start_index}, {t2_end_index}")
    #     return 0, "", ""

    t1 = (full_text_list[t1_start_index:t1_end_index])
    t2 = (full_text_list[t2_start_index:t2_end_index])
    free_attempt = (full_text_list[t1_start_index:free_attempt_end_index])

    t1_processed = list(
        filter(lambda x: not ((x == "") or ("speaker 0" in x.lower()) or (x is None) or ("first prompt" in x.lower())),
               t1))
    t2_processed = list(
        filter(lambda x: not ((x == "") or ("speaker 0" in x.lower()) or (x is None) or ("first prompt" in x.lower())),
               t2))
    free_attempt_processed = list(
        filter(lambda x: not ((x == "") or ("speaker 0" in x.lower()) or (x is None) or ("first prompt" in x.lower())),
               free_attempt))

    for i in range(0, len(t1_processed)):
        t1_processed[i] = t1_processed[i].replace("Speaker 1:", "")
        t1_processed[i] = t1_processed[i].replace("(?)", "")
        # t1_processed[i].strip()

    for i in range(0, len(t2_processed)):
        t2_processed[i] = t2_processed[i].replace("Speaker 1:", "")
        t2_processed[i] = t2_processed[i].replace("(?)", "")
        # t2_processed[i].strip()

    for i in range(0, len(free_attempt_processed)):
        free_attempt_processed[i] = free_attempt_processed[i].replace("Speaker 1:", "")
        free_attempt_processed[i] = free_attempt_processed[i].replace("(?)", "")
        # free_attempt_processed[i].strip()

    t1_final = re.sub(r'\s+', ' ', "".join(t1_processed))
    t1_final = t1_final.replace("\n", " ")
    t2_final = re.sub(r'\s+', ' ', "".join(t2_processed))
    t2_final = t2_final.replace("\n", " ")
    free_attempt_final = re.sub(r'\s+', ' ', "".join(free_attempt_processed))
    free_attempt_final = free_attempt_final.replace("\n", " ")
    return t1_final, t2_final, free_attempt_final


def read_and_compute(file_path, baseline):
    full_text_list = read_docx(file_path)

    t1_start_index = 0
    t1_end_index = 0
    print(full_text_list)
    for paragraph in full_text_list:
        if ("storytelling" in paragraph.lower()) or ("mirar la imagen" in paragraph.lower()) or (
                "mira la imagen" in paragraph.lower()) or ("observa la imagen" in paragraph.lower()) or (
                "mires la imagen" in paragraph.lower()) or ("la imagen de esta pantalla" in paragraph.lower()) or (
                "una imagen en la pantalla" in paragraph.lower()) or ("Mirarás la imagen en " in paragraph):
            break
        # This is for P16_S1, it doesn't have storytelling keyword.
        if " You will look at the screen. Um you will see a picture" in paragraph:
            break
        t1_start_index += 1
    for paragraph in full_text_list:
        if "t2" in paragraph.lower():
            break
        t1_end_index += 1

    t2_start_index = 0
    t2_end_index = len(full_text_list) - 1
    for paragraph in full_text_list:
        if ("different game" in paragraph.lower()) or ("juego diferente" in paragraph.lower()) or (
                "animal" in paragraph.lower()) or ("frutas" in paragraph.lower()):
            break
        if "different challenge" in paragraph.lower():
            break
        t2_start_index += 1

    # if t1_end_index > t2_start_index or t2_start_index > t2_end_index:
    #     print(f"Error in file: {file_path}")
    #     print(f"Info: {t1_start_index}, {t1_end_index}, {t2_start_index}, {t2_end_index}")
    #     return 0, "", ""

    t1 = (full_text_list[t1_start_index:t1_end_index])
    t2 = (full_text_list[t2_start_index:t2_end_index])

    t1_processed = list(
        filter(lambda x: not ((x == "") or ("speaker 0" in x.lower()) or (x is None) or ("first prompt" in x.lower())),
               t1))
    t2_processed = list(
        filter(lambda x: not ((x == "") or ("speaker 0" in x.lower()) or (x is None) or ("first prompt" in x.lower())),
               t2))

    for i in range(0, len(t1_processed)):
        t1_processed[i] = t1_processed[i].replace("Speaker 1:", "")
        t1_processed[i] = t1_processed[i].replace("(?)", "")
        t1_processed[i].strip()

    for i in range(0, len(t2_processed)):
        t2_processed[i] = t2_processed[i].replace("Speaker 1:", "")
        t2_processed[i] = t2_processed[i].replace("(?)", "")
        t2_processed[i].strip()

    test_text = re.sub(r'\s+', ' ', "".join(t1_processed))

    baseline_text_emb = get_embedding(baseline)
    try:
        test_text_emb = get_embedding(test_text)
    except Exception as e:
        print(f"Error: {e}")
        print(f"Error in file: {file_path}")
        print(f"Error in text: {test_text}")
        print(f"Info: {t1_start_index}, {t1_end_index}, {t2_start_index}, {t2_end_index}")
        return 0, "".join(t1_processed), "".join(t2_processed)
    baseline_text_emb, test_text_emb = pad_embeddings(baseline_text_emb, test_text_emb)

    return cosine_similarity(baseline_text_emb, test_text_emb), re.sub(r'\s+', ' ', "".join(t1_processed)), re.sub(
        r'\s+', ' ', "".join(t2_processed))


lang = "ES"

directory_path = f"./data/{lang}/"
output_path = f"./data_processed/{lang}/"
t1_output_path = f"./data_processed/{lang}/Original/T1/"
t2_output_path = f"./data_processed/{lang}/Original/T2/"
free_attempt_output_path = f"./data_processed/{lang}/Original/Free_Attempt/"
t1_no_pun_output_path = f"./data_processed/{lang}/NoPunctuation/T1/"
t2_no_pun_output_path = f"./data_processed/{lang}/NoPunctuation/T2/"
free_attempt_no_pun_output_path = f"./data_processed/{lang}/NoPunctuation/Free_Attempt/"

if lang == "EN":
    baseline_picnic_text = baseline_picnic_text_en
    baseline_cookie_text = baseline_cookie_text_en
else:
    baseline_picnic_text = baseline_picnic_text_es
    baseline_cookie_text = baseline_cookie_text_es

# Get all file names in the directory
file_names = [f for f in os.listdir(directory_path) if os.path.isfile(os.path.join(directory_path, f))]

with open(f"{output_path}similarity.txt", "w") as f:
    f.write("Similarity Score: \n")

# Print all file names
for file_name in file_names:
    print(f"Processing: {file_name}")
    base_name = file_name.split(".")[0]
    base_name.replace("_check", "")
    if "S2" in base_name or "S4" in base_name:
        baseline_text = baseline_picnic_text
    else:
        baseline_text = baseline_cookie_text
    # similarity, t1, t2 = read_and_compute(f"{directory_path}{file_name}", baseline_text)
    t1, t2, free_attempt = read_and_preprocess_text(f"{directory_path}{file_name}")
    if lang == "EN":
        t1 = t1.replace("í", "'")
        t2 = t2.replace("í", "'")
        free_attempt = free_attempt.replace("í", "'")
    # print(f"{base_name} Similarity: {similarity:.2f}")
    # with open(f"{output_path}similarity.txt", "a") as f:
    #     f.write(f"{base_name}: {similarity:.2f}\n")

    print(f"Writing processed T1 text to file: {t1_output_path}{base_name}_t1.txt")
    with open(f"{t1_output_path}{base_name}_t1.txt", "w") as f:
        f.write(t1)
    with open(f"{t1_no_pun_output_path}{base_name}_no_pun_t1.txt", "w") as f:
        f.write(remove_punctuation(t1))
    # print(remove_punctuation(t1))

    print(f"Writing processed T2 text to file: {t2_output_path}{base_name}_t2.txt")
    with open(f"{t2_output_path}{base_name}_t2.txt", "w") as f:
        f.write(t2)
    with open(f"{t2_no_pun_output_path}{base_name}_no_pun_t2.txt", "w") as f:
        f.write(remove_punctuation(t2))
    # print(remove_punctuation(t2))

    print(f"Writing processed Free Attempt text to file: {free_attempt_output_path}{base_name}_free_attempt.txt")
    with open(f"{free_attempt_output_path}{base_name}_free_attempt.txt", "w") as f:
        f.write(t2)
    with open(f"{free_attempt_no_pun_output_path}{base_name}_no_pun_free_attempt.txt", "w") as f:
        f.write(remove_punctuation(free_attempt))
